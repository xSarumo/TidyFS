from pathlib import Path
from pypdf import PdfReader
from docx import Document
import logging

logging.getLogger("pypdf").setLevel(logging.ERROR)

PROJECT_ROOT = Path(__file__).resolve().parent.parent
FILES_FILE = PROJECT_ROOT / "files" / "files.json"

PREVIEW_CHARS = 4000

def extract_preview(path: str, ext: str) -> str:
    ext = ext.lower()

    if ext == ".pdf":
        return extract_pdf(path)
    elif ext in [".docx", "doc"]:
        return extract_doc(path)
    
    return ""

def limit_text(text: str) -> str:
    return text[:PREVIEW_CHARS]

def extract_pdf(path: str) -> str:
    try:
        reader = PdfReader(path, strict=False)

        chunks = []

        for page in reader.pages[:2]:
            text = page.extract_text() or ""
            chunks.append(text)

            if len("\n".join(chunks)) >= PREVIEW_CHARS:
                break
        return limit_text("\n".join(chunks))
    except Exception:
        return ""
    

def extract_doc(path: str) -> str:
    try:
        document = Document(path)

        chunks = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                chunks.append(text)

            if len("\n".join(chunks)) >= PREVIEW_CHARS:
                return limit_text("\n".join(chunks))
            

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        chunks.append(text)

                    if len("\n".join(chunks)) >= PREVIEW_CHARS:
                        return limit_text("\n".join(chunks))

        return limit_text("\n".join(chunks))
    except Exception:
        return ""